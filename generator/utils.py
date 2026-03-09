from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from io import BytesIO
import PyPDF2
import docx

def generate_pdf_file(questions, topic, include_answers=True):
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=50, rightMargin=50, topMargin=50, bottomMargin=50)
        
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#00d4ff'),
            spaceAfter=30,
            alignment=0
        )
        
        question_style = ParagraphStyle(
            'CustomQuestion',
            parent=styles['BodyText'],
            fontSize=11,
            textColor=colors.black,
            spaceAfter=6,
            leading=14,
            fontName='Helvetica-Bold'
        )
        
        answer_style = ParagraphStyle(
            'CustomAnswer',
            parent=styles['Italic'],
            fontSize=10,
            textColor=colors.HexColor('#2d5016'),
            spaceAfter=6,
            leading=13,
            leftIndent=20,
            fontName='Helvetica'
        )
        
        explanation_style = ParagraphStyle(
            'CustomExplanation',
            parent=styles['BodyText'],
            fontSize=9,
            textColor=colors.HexColor('#1a4d7c'),
            spaceAfter=12,
            leading=12,
            leftIndent=20,
            fontName='Helvetica',
            backColor=colors.HexColor('#f0f7ff'),
            borderPadding=8
        )
        
        step_style = ParagraphStyle(
            'StepStyle',
            parent=styles['BodyText'],
            fontSize=9,
            textColor=colors.HexColor('#333333'),
            spaceAfter=3,
            leading=11,
            leftIndent=30,
            fontName='Helvetica'
        )
        
        story = []
        
        # Add title
        suffix = "" if include_answers else " (Questions Only)"
        title_text = f"Generated Questions: {topic}{suffix}"
        title_p = Paragraph(title_text, title_style)
        story.append(title_p)
        story.append(Spacer(1, 0.3*inch))
        
        # Add questions
        for i, q in enumerate(questions, 1):
            # Question with type indicator for numerical problems
            q_type = getattr(q, 'question_type', 'mixed')
            type_label = " [Numerical]" if q_type == 'numerical' else ""
            question_text = f"<b>{i}. {q.text}{type_label}</b>"
            question_p = Paragraph(question_text, question_style)
            story.append(question_p)
            story.append(Spacer(1, 0.08*inch))
            
            if include_answers:
                # Answer - with label
                if q.answer:
                    answer_text = f"<b>Answer:</b> {q.answer}"
                    answer_p = Paragraph(answer_text, answer_style)
                    story.append(answer_p)
                
                # Step-by-step explanation for numerical problems
                explanation = getattr(q, 'explanation', '')
                if explanation:
                    story.append(Spacer(1, 0.05*inch))
                    explanation_header = Paragraph("<b>Step-by-Step Solution:</b>", explanation_style)
                    story.append(explanation_header)
                    
                    # Parse and format steps
                    steps = explanation.replace('Step ', '\nStep ').strip().split('\n')
                    for step in steps:
                        step = step.strip()
                        if step:
                            step_p = Paragraph(step, step_style)
                            story.append(step_p)
            
            story.append(Spacer(1, 0.25*inch))
            
            # Add page break after every few questions
            if include_answers:
                questions_per_page = 2 if q_type == 'numerical' else 3
            else:
                questions_per_page = 5
            if (i % questions_per_page == 0) and (i < len(questions)):
                story.append(PageBreak())
        
        # Build the document
        doc.build(story)
        
        buffer.seek(0)
        print(f"DEBUG: PDF buffer size: {len(buffer.getvalue())} bytes")
        return buffer
        
    except Exception as e:
        print(f"ERROR in generate_pdf_file: {str(e)}")
        import traceback
        traceback.print_exc()
        raise


def extract_text_from_file(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith('.pdf'):
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    elif name.endswith('.docx'):
        document = docx.Document(uploaded_file)
        text = "\n".join(para.text for para in document.paragraphs)
        return text.strip()
    elif name.endswith('.txt'):
        raw = uploaded_file.read()
        if isinstance(raw, bytes):
            return raw.decode('utf-8').strip()
        return raw.strip()
    else:
        raise ValueError(f"Unsupported file type: {name}")


def generate_professional_pdf(questions, topic, include_answers=True, institution="", duration="", total_marks=None):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=50, rightMargin=50, topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()

    institution_style = ParagraphStyle(
        'Institution',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.black,
        alignment=1,
        spaceAfter=4,
        fontName='Helvetica-Bold'
    )

    exam_title_style = ParagraphStyle(
        'ExamTitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.black,
        alignment=1,
        spaceAfter=4,
        fontName='Helvetica-Bold'
    )

    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['BodyText'],
        fontSize=10,
        textColor=colors.black,
        alignment=1,
        spaceAfter=2,
        fontName='Helvetica'
    )

    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=colors.black,
        spaceAfter=8,
        spaceBefore=16,
        fontName='Helvetica-Bold',
        underlineProportion=1
    )

    q_style = ParagraphStyle(
        'ProfQuestion',
        parent=styles['BodyText'],
        fontSize=11,
        textColor=colors.black,
        spaceAfter=4,
        leading=14,
        fontName='Helvetica'
    )

    option_style = ParagraphStyle(
        'Option',
        parent=styles['BodyText'],
        fontSize=10,
        textColor=colors.black,
        spaceAfter=2,
        leading=13,
        leftIndent=20,
        fontName='Helvetica'
    )

    answer_key_title_style = ParagraphStyle(
        'AnswerKeyTitle',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.black,
        alignment=1,
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )

    answer_style = ParagraphStyle(
        'ProfAnswer',
        parent=styles['BodyText'],
        fontSize=10,
        textColor=colors.HexColor('#2d5016'),
        spaceAfter=4,
        leading=13,
        fontName='Helvetica'
    )

    story = []

    # Header
    if institution:
        story.append(Paragraph(institution, institution_style))
    story.append(Paragraph(f"Examination: {topic}", exam_title_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=6))

    meta_parts = []
    if duration:
        meta_parts.append(f"Duration: {duration}")
    if total_marks is not None:
        meta_parts.append(f"Total Marks: {total_marks}")
    if meta_parts:
        story.append(Paragraph(" | ".join(meta_parts), meta_style))
    story.append(Spacer(1, 0.2 * inch))

    # Organize questions by type into sections
    section_map = {
        'mcq': 'Section A: Multiple Choice Questions',
        'short': 'Section B: Short Answer Questions',
        'long': 'Section C: Long Answer Questions',
        'numerical': 'Section D: Numerical Problems',
        'true_false': 'Section E: True or False',
    }
    ordered_types = ['mcq', 'short', 'long', 'numerical', 'true_false']

    sections = {}
    for q in questions:
        q_type = getattr(q, 'question_type', 'short')
        sections.setdefault(q_type, []).append(q)

    global_num = 1
    answer_entries = []

    for q_type in ordered_types:
        if q_type not in sections:
            continue
        section_title = section_map.get(q_type, f"Section: {q_type.replace('_', ' ').title()}")
        story.append(Paragraph(section_title, section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceAfter=8))

        for q in sections[q_type]:
            marks = getattr(q, 'marks', None)
            marks_text = f"  [{marks} marks]" if marks else ""
            story.append(Paragraph(f"{global_num}. {q.text}{marks_text}", q_style))

            # MCQ options
            if q_type == 'mcq':
                options = getattr(q, 'options', None)
                if options and isinstance(options, list):
                    labels = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h']
                    for idx, opt in enumerate(options):
                        label = labels[idx] if idx < len(labels) else str(idx + 1)
                        story.append(Paragraph(f"({label}) {opt}", option_style))

            story.append(Spacer(1, 0.12 * inch))

            if include_answers and getattr(q, 'answer', None):
                explanation = getattr(q, 'explanation', '')
                answer_entries.append((global_num, q.answer, explanation))

            global_num += 1

    # Also handle any question types not in ordered_types
    for q_type, qs in sections.items():
        if q_type in ordered_types:
            continue
        section_title = f"Section: {q_type.replace('_', ' ').title()}"
        story.append(Paragraph(section_title, section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceAfter=8))

        for q in qs:
            marks = getattr(q, 'marks', None)
            marks_text = f"  [{marks} marks]" if marks else ""
            story.append(Paragraph(f"{global_num}. {q.text}{marks_text}", q_style))
            story.append(Spacer(1, 0.12 * inch))

            if include_answers and getattr(q, 'answer', None):
                explanation = getattr(q, 'explanation', '')
                answer_entries.append((global_num, q.answer, explanation))

            global_num += 1

    # Answer Key
    if include_answers and answer_entries:
        story.append(PageBreak())
        story.append(Paragraph("Answer Key", answer_key_title_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=10))

        for num, answer, explanation in answer_entries:
            story.append(Paragraph(f"<b>{num}.</b> {answer}", answer_style))
            if explanation:
                story.append(Paragraph(f"<i>Explanation: {explanation}</i>", answer_style))
            story.append(Spacer(1, 0.06 * inch))

    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_docx_file(questions, topic, include_answers=True):
    document = docx.Document()

    # Title
    title = document.add_heading(f"Questions: {topic}", level=0)

    document.add_paragraph("")

    for i, q in enumerate(questions, 1):
        marks = getattr(q, 'marks', None)
        marks_text = f"  [{marks} marks]" if marks else ""
        q_para = document.add_paragraph()
        run = q_para.add_run(f"{i}. {q.text}{marks_text}")
        run.bold = True

        # MCQ options
        q_type = getattr(q, 'question_type', '')
        if q_type == 'mcq':
            options = getattr(q, 'options', None)
            if options and isinstance(options, list):
                labels = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h']
                for idx, opt in enumerate(options):
                    label = labels[idx] if idx < len(labels) else str(idx + 1)
                    document.add_paragraph(f"  ({label}) {opt}")

    # Answer section
    if include_answers:
        document.add_page_break()
        document.add_heading("Answer Key", level=1)

        for i, q in enumerate(questions, 1):
            answer = getattr(q, 'answer', None)
            if answer:
                ans_para = document.add_paragraph()
                ans_para.add_run(f"{i}. ").bold = True
                ans_para.add_run(str(answer))

                explanation = getattr(q, 'explanation', '')
                if explanation:
                    exp_para = document.add_paragraph()
                    run = exp_para.add_run(f"   Explanation: {explanation}")
                    run.italic = True

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer